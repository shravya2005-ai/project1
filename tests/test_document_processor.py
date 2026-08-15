import tempfile
import unittest
from pathlib import Path
import docx
from src.document_processor import extract_texts_from_file, _split_text_into_chunks


class TestDocumentProcessor(unittest.TestCase):

    def test_split_text_into_chunks(self):
        text = "Paragraph 1 sentence 1. Paragraph 1 sentence 2.\n\nParagraph 2 sentence 1. Paragraph 2 sentence 2."
        chunks = _split_text_into_chunks(text, chunk_size=50, chunk_overlap=10)
        self.assertGreaterEqual(len(chunks), 1)
        self.assertTrue(any("Paragraph 1" in c for c in chunks))

    def test_extract_texts_from_txt(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "sample.txt"
            file_path.write_text("This is a test document. It contains a few sentences for testing.")

            chunks = extract_texts_from_file(file_path)
            self.assertGreater(len(chunks), 0)
            self.assertEqual(chunks[0]["metadata"]["source"], "sample.txt")
            self.assertIn("chunk_id", chunks[0]["metadata"])
            self.assertIn("This is a test document", chunks[0]["document"])

    def test_extract_texts_from_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "sample.docx"
            doc = docx.Document()
            doc.add_paragraph("First paragraph of docx test.")
            doc.add_paragraph("Second paragraph of docx test.")
            doc.save(file_path)

            chunks = extract_texts_from_file(file_path)
            self.assertGreater(len(chunks), 0)
            self.assertEqual(chunks[0]["metadata"]["source"], "sample.docx")
            self.assertIn("First paragraph", chunks[0]["document"])


if __name__ == "__main__":
    unittest.main()
